from docx import Document

doc = Document()
doc.add_heading("Test Document", level=1)
doc.add_paragraph("This is a sample paragraph for testing Word to PDF conversion.")

doc.save("test.docx")
