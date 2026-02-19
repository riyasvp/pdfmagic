#!/usr/bin/env python3
"""
OCR PDF - Extract text from scanned PDFs using pytesseract.
Usage: python ocr_pdf.py <input_pdf> [language]
Output: JSON with result
"""

import sys
import os
import json
from datetime import datetime

try:
    import pytesseract
    from pdf2image import convert_from_path
    from docx import Document
    from docx.shared import Pt
except ImportError as e:
    print(json.dumps({"success": False, "error": f"Missing dependency: {str(e)}"}))
    sys.exit(1)

DOWNLOAD_DIR = "/home/z/my-project/download"

def ocr_pdf(input_path, language="eng"):
    """Perform OCR on PDF and extract text."""
    if not os.path.exists(input_path):
        return {"success": False, "error": f"File not found: {input_path}"}
    
    try:
        # Convert PDF to images
        images = convert_from_path(input_path, dpi=300)
        
        if not images:
            return {"success": False, "error": "No pages found in PDF"}
        
        # Create Word document for output
        doc = Document()
        title = doc.add_heading('OCR Extracted Text', 0)
        
        # Process each page
        for i, image in enumerate(images):
            # Add page header
            doc.add_heading(f'Page {i + 1}', level=2)
            
            # Perform OCR
            text = pytesseract.image_to_string(image, lang=language)
            
            if text.strip():
                # Add extracted text
                for para in text.split('\n\n'):
                    if para.strip():
                        p = doc.add_paragraph(para.strip())
                        p.style.font.size = Pt(11)
            
            doc.add_paragraph()  # Add spacing between pages
        
        # Generate output filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"ocr_result_{timestamp}.docx"
        output_path = os.path.join(DOWNLOAD_DIR, output_filename)
        
        # Ensure download directory exists
        os.makedirs(DOWNLOAD_DIR, exist_ok=True)
        
        # Save document
        doc.save(output_path)
        
        return {
            "success": True,
            "output": output_path,
            "pages_processed": len(images)
        }
    
    except Exception as e:
        return {"success": False, "error": str(e)}

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"success": False, "error": "Input PDF file required"}))
        sys.exit(1)
    
    input_path = sys.argv[1]
    language = sys.argv[2] if len(sys.argv) > 2 else "eng"
    result = ocr_pdf(input_path, language)
    print(json.dumps(result))
