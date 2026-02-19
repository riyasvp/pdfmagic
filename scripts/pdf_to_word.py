#!/usr/bin/env python3
"""
Convert PDF to Word document.
Usage: python pdf_to_word.py <input_pdf>
Output: JSON with result
"""

import sys
import os
import json
from datetime import datetime

try:
    import pdfplumber
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as e:
    print(json.dumps({"success": False, "error": f"Missing dependency: {str(e)}"}))
    sys.exit(1)

DOWNLOAD_DIR = "/home/z/my-project/download"

def extract_text_with_formatting(pdf_path):
    """Extract text from PDF with basic formatting preservation."""
    paragraphs = []
    
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                # Split into paragraphs
                for para in text.split('\n\n'):
                    if para.strip():
                        paragraphs.append(para.strip())
    
    return paragraphs

def convert_pdf_to_word(input_path):
    """Convert PDF to Word document."""
    if not os.path.exists(input_path):
        return {"success": False, "error": f"File not found: {input_path}"}
    
    try:
        # Extract text from PDF
        paragraphs = extract_text_with_formatting(input_path)
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Converted PDF Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add paragraphs
        for para_text in paragraphs:
            p = doc.add_paragraph(para_text)
            p.style.font.size = Pt(11)
        
        # Generate output filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"converted_{timestamp}.docx"
        output_path = os.path.join(DOWNLOAD_DIR, output_filename)
        
        # Ensure download directory exists
        os.makedirs(DOWNLOAD_DIR, exist_ok=True)
        
        # Save document
        doc.save(output_path)
        
        return {"success": True, "output": output_path}
    
    except Exception as e:
        return {"success": False, "error": str(e)}

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"success": False, "error": "Input PDF file required"}))
        sys.exit(1)
    
    input_path = sys.argv[1]
    result = convert_pdf_to_word(input_path)
    print(json.dumps(result))
